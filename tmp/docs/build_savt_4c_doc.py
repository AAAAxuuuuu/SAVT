from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(r"G:\SAVT")
DOWNLOADS = Path.home() / "Downloads"
TEMPLATE_DIR = next(
    d for d in DOWNLOADS.iterdir()
    if d.is_dir() and "2023" in d.name and any(f.name.startswith("01-3") for f in d.iterdir())
)
TEMPLATE = next(f for f in TEMPLATE_DIR.iterdir() if f.name.startswith("01-3") and f.suffix.lower() == ".docx")
OUTPUT_DIR = ROOT / "output" / "doc"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT = OUTPUT_DIR / "SAVT-软件应用与开发类作品设计和开发文档-2026-04-13.docx"


def clear_document(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def set_run_font(run, size=11, bold=False, font="宋体"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold


def add_para(doc, text="", size=11, bold=False, align=None, first_line=True):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if first_line:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def add_title(doc, text, size=18, bold=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, font="黑体")
    return p


def add_h1(doc, text):
    p = doc.add_heading(level=1)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=15, bold=True, font="黑体")
    return p


def add_sub(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run_font(run, size=12, bold=True, font="黑体")
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing = 1.2
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(item)
        set_run_font(run, size=10.5)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_run_font(run, size=10, bold=True, font="黑体")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.line_spacing = 1.1
            run = p.add_run(value)
            set_run_font(run, size=9.5)
    doc.add_paragraph()
    return table


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    set_run_font(run, size=9.5, font="Consolas")


def build():
    doc = Document(TEMPLATE)
    clear_document(doc)
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)

    add_title(doc, "中国大学生计算机设计大赛", size=20)
    add_title(doc, "软件应用与开发类作品设计和开发文档", size=18)
    for _ in range(2):
        doc.add_paragraph()
    fields = [
        ("作品编号", "待报名系统生成后填写"),
        ("作品名称", "SAVT 软件架构分析与可视化系统"),
        ("作者", "待参赛团队确认后填写"),
        ("版本编号", "v0.1.0-competition-draft, Git 87828d6"),
        ("填写日期", "2026年4月13日"),
    ]
    for label, value in fields:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label}：{value}")
        set_run_font(run, size=13)
        p.paragraph_format.space_after = Pt(8)
    doc.add_paragraph()
    add_para(
        doc,
        "真实性说明：本文档依据当前本地仓库 G:\\SAVT、Git 提交 87828d6f6dd974ed92f8fa93525debd65af4ad7a、"
        "本轮 Windows MSVC 构建与 ctest 实测结果、项目内阶段记录以及公开可核验的标准和论文资料撰写。"
        "作品编号、作者、学校、指导教师、最终截图、演示视频和正式报名材料需由参赛团队按报名系统信息补齐，本文不编造上述信息。",
        size=10.5,
    )
    doc.add_page_break()

    add_h1(doc, "需求分析")
    add_para(
        doc,
        "SAVT 面向中大型软件项目的架构理解问题。随着项目规模增长，开发者面对的不再是单个源文件，"
        "而是跨目录、跨模块、跨语言、跨构建环境的复杂工程。传统文件树和全文检索能够定位局部代码，"
        "但难以回答系统入口在哪里、核心能力如何分布、模块之间为什么存在依赖、当前分析结论是否具备语义精度等问题。"
        "软件架构重构研究指出，架构并不总是显式存在于代码中，系统演进还会造成架构漂移，因此从源码中重建和校验架构是维护理解的重要辅助任务。",
    )
    add_para(
        doc,
        "本作品的目标用户包括软件工程课程团队、开源项目维护者、接手陌生仓库的新成员、架构评审人员和重构负责人。"
        "他们需要一种工具把源码事实转化为可交互、可解释、可追溯的架构视图，并在精度不足时明确告诉用户限制原因。"
        "SAVT 的核心需求可以概括为五点：项目扫描、源码解析、架构聚合、图形化展示、可信解释。"
    )
    add_table(
        doc,
        ["对比对象", "优势", "局限", "SAVT 的定位"],
        [
            ["IDE 文件树与跳转", "定位定义和调用方便", "偏局部阅读，难以形成全局架构视图", "生成系统上下文、能力图和组件图"],
            ["文档生成器", "可批量生成 API 文档", "文档多围绕接口，缺少精度状态和证据链", "把模块、关系、证据和报告统一组织"],
            ["静态分析工具", "可做较深入的语法或语义检查", "输出常偏工程人员，架构表达和交互弱", "面向架构理解和答辩展示"],
            ["通用大模型", "自然语言解释能力强", "可能脱离证据产生幻觉，且隐私边界不清", "只允许在受限证据包内辅助解释"],
        ],
    )
    add_para(
        doc,
        "需求边界方面，SAVT 不承诺在缺少编译数据库时完成工业级 C++ 语义分析；在没有网络和 API 配置时，AI 解释功能可以不可用，"
        "但项目扫描、语法分析、架构图、组件下钻、AST 预览和报告生成仍应独立运行。"
    )

    add_h1(doc, "概要设计")
    add_para(
        doc,
        "SAVT 采用分层架构，核心原则是分析事实与界面渲染解耦。底层负责项目扫描和源码解析，中间层负责统一图模型、能力图、组件图和证据包，"
        "上层负责布局计算、Qt 服务编排、QML 场景映射和桌面交互展示。该设计与 ISO/IEC/IEEE 42010 中关于架构描述、视图和模型的思想一致，"
        "即一个架构描述应围绕利益相关者关切组织多个视图，而不是只给出单一图形。"
    )
    add_code_block(
        doc,
        "QML Desktop Workbench\n"
        "  -> AnalysisController / AnalysisOrchestrator\n"
        "  -> IncrementalAnalysisPipeline\n"
        "  -> CppProjectAnalyzer\n"
        "  -> SyntaxProjectAnalyzer or SemanticProjectAnalyzer\n"
        "  -> AnalysisReport\n"
        "  -> ArchitectureOverview / CapabilityGraph / ComponentGraph\n"
        "  -> LayeredGraphLayout\n"
        "  -> SceneMapper\n"
        "  -> L1-L4 Views, Report, AST Preview, AI Explanation",
    )
    add_table(
        doc,
        ["模块", "路径", "主要职责"],
        [
            ["桌面入口", "apps/desktop", "创建 Qt 应用、注入 AnalysisController、加载 QML 工作台"],
            ["核心模型", "src/core", "定义 AnalysisReport、ArchitectureOverview、CapabilityGraph、ComponentGraph 和证据包"],
            ["解析适配", "src/parser", "封装 Tree-sitter C++ 解析能力"],
            ["分析后端", "src/analyzer", "扫描项目、抽取符号和关系、管理语法/语义精度模式"],
            ["布局算法", "src/layout", "计算能力图、组件图和模块图布局"],
            ["UI 服务", "src/ui", "异步分析、缓存、报告、AST 预览、语义状态、AI 服务和 scene 映射"],
            ["AI 适配", "src/ai", "读取 DeepSeek/OpenAI 兼容配置、构建请求、解析结构化响应"],
            ["测试体系", "tests", "后端、快照、AI、UI 服务层、性能基线和 UTF-8 检查"],
        ],
    )
    add_para(
        doc,
        "系统数据流从用户选择项目目录开始。分析器先生成扫描清单和元数据指纹，再根据精度模式选择 Tree-sitter 语法路径或 libclang 语义路径。"
        "完成分析后，结果被聚合为架构总览和能力图，再进入布局模块生成可视化坐标，最终由 QML 工作台展示。"
        "项目不使用业务数据库，主要状态来自内存模型、项目级 JSON 配置和可选本地 AI 配置文件。"
    )

    add_h1(doc, "详细设计")
    add_sub(doc, "1. 精度模式与语义诊断")
    add_para(
        doc,
        "SAVT 定义 SyntaxOnly、SemanticPreferred 和 SemanticRequired 三种分析模式。SyntaxOnly 只执行语法级分析；SemanticPreferred 优先尝试语义分析，"
        "失败时可回退到语法分析；SemanticRequired 必须进入语义分析，否则直接生成阻断报告。阻断原因包括 missing_compile_commands、backend_unavailable、"
        "llvm_not_found、llvm_headers_missing、libclang_not_found、system_headers_unresolved 和 translation_unit_parse_failed 等。"
    )
    add_code_block(
        doc,
        "if precision requests semantics:\n"
        "    locate compile_commands.json\n"
        "    inspect LLVM/libclang availability\n"
        "    if compilation database and backend are ready:\n"
        "        run semantic analyzer\n"
        "        return semantic report on success\n"
        "    if precision == SemanticRequired:\n"
        "        return blocked report with exact reason\n"
        "run Tree-sitter syntax analyzer\n"
        "return syntax or syntax_fallback report",
    )
    add_sub(doc, "2. 图模型与证据包")
    add_para(
        doc,
        "底层 AnalysisReport 保存符号节点、关系边、解析文件数、语义状态码和诊断信息。CapabilityGraph 将模块组织为入口、能力和基础设施节点，"
        "并在每个节点上保留事实、规则、结论、来源文件、符号、模块、关系、置信标签和置信原因。该证据包机制使评委或开发者可以追问"
        "“为什么存在这个节点”以及“这条边来自什么证据”，避免架构图变成不可审查的静态图片。"
    )
    add_sub(doc, "3. 增量分析与缓存")
    add_para(
        doc,
        "增量管线把一次分析拆分为扫描、解析、聚合、布局四层缓存。缓存键绑定项目根路径、分析选项、项目配置签名、编译数据库签名、源码元数据指纹和缓存版本。"
        "当源码、项目配置、compile_commands.json 或精度模式发生变化时，对应层及其下游层会失效重算。"
    )
    add_table(
        doc,
        ["缓存层", "缓存内容", "失效条件"],
        [
            ["扫描层", "项目源码清单和元数据指纹", "文件增删、大小或时间戳变化，分析选项变化"],
            ["解析层", "AnalysisReport", "扫描指纹、编译数据库、精度模式或项目配置变化"],
            ["聚合层", "ArchitectureOverview、CapabilityGraph、规则报告", "解析层结果变化"],
            ["布局层", "L2/L3 scene layout 和组件图", "聚合层结果变化"],
        ],
    )
    add_sub(doc, "4. 人机界面与典型流程")
    add_para(
        doc,
        "当前桌面主界面由 Qt Quick/QML 实现，主入口为 apps/desktop/qml/Main.qml 和 greenfield 工作台。用户可以选择项目目录、启动分析、停止分析、查看语义状态、"
        "在架构全景图中缩放和平移、选择节点、聚焦上下游关系、下钻组件图、查看 AST 预览，并复制 Markdown 报告。"
        "最终参赛版建议补充实际界面截图和 10 分钟演示视频，本稿不伪造截图。"
    )
    add_sub(doc, "5. AI 辅助解释边界")
    add_para(
        doc,
        "AI 功能是可选辅助层。系统只在用户手动触发时读取受限证据包，构造 DeepSeek 或 OpenAI 兼容 Chat Completions 请求。"
        "内置规则要求模型只处理 SAVT 架构阅读任务，只使用提供的证据，不臆造文件、模块、符号或业务含义，并以结构化 JSON 返回。"
        "无网络或无本地密钥配置时，不影响核心静态分析功能。"
    )

    add_h1(doc, "测试报告")
    add_para(
        doc,
        "本节数据为 2026年4月13日 在当前 Windows 工作区实测结果。构建命令在 Visual Studio 2022 Developer Command Prompt x64 环境下执行，"
        "目标 savt_backend_tests、savt_snapshot_tests、savt_ai_tests、savt_ui_tests、savt_perf_baseline 已构建成功。随后执行 ctest 全量测试。"
    )
    add_table(
        doc,
        ["项目", "实测结果"],
        [
            ["Git 提交", "87828d6f6dd974ed92f8fa93525debd65af4ad7a"],
            ["统计范围", "src、apps/desktop、tests/cpp、config、scripts、samples、docs；排除 build、third_party、docs/competition"],
            ["文件规模", "179 个文件；其中 QML 68、C++ 源文件 35、头文件 29、Markdown 25"],
            ["主要代码行数", "C/C++ 相关 63 文件 19165 行；QML/JS 70 文件 12200 行；C++ 测试 5 文件 2674 行"],
            ["精度样例", "tests/fixtures/precision 21 个；tests/golden/precision 21 个"],
            ["技术债标记扫描", "rg TODO/FIXME/HACK/XXX/@deprecated 未发现匹配项"],
        ],
    )
    add_table(
        doc,
        ["测试项", "本轮结果", "说明"],
        [
            ["savt_ai_tests", "通过", "AI 配置、请求构造和响应解析相关测试通过"],
            ["savt_perf_smoke", "通过", "性能 smoke 输出了冷/热缓存耗时数据"],
            ["savt_utf8_check", "通过", "仓库文本 UTF-8 字节检查通过"],
            ["savt_backend_tests", "失败", "失败断言：entry overview modules should become entry components in the L3 graph"],
            ["savt_snapshot_tests", "失败", "多个 fixture 与 golden 快照不一致，说明当前实现已改变输出，需要审查并更新基线或修复回归"],
            ["savt_ui_tests", "失败", "失败断言：system context cards should publish the expected summary cards"],
        ],
    )
    add_para(
        doc,
        "因此，本文档不将当前仓库描述为“全部测试通过”的最终交付状态。更准确的结论是：项目可构建，AI、性能 smoke 和编码检查通过；"
        "后端 L3 组件映射、系统上下文卡片和快照基线存在待修复或待重新审定项。正式参赛前必须修复失败测试，或对 golden 基线变化给出代码审查记录。"
    )
    add_table(
        doc,
        ["性能样例", "发现文件", "解析文件", "能力节点/边", "冷启动总耗时(ms)", "热缓存总耗时(ms)", "缓存命中", "峰值内存(MB)"],
        [
            ["cpp_vendored_dependency_app", "6", "6", "2/1", "19.92", "3.02", "四层命中", "18.23"],
            ["qml_dashboard_workspace", "8", "5", "5/2", "18.33", "6.17", "四层命中", "18.47"],
            ["synthetic_quick_2000", "2001", "2001", "9/0", "2350.33", "469.15", "四层命中", "77.93"],
        ],
    )
    add_para(
        doc,
        "上述性能数据来自 savt_perf_smoke 的本轮实测输出。由于测试机器负载、Debug 构建、磁盘状态和缓存状态会影响绝对耗时，"
        "该结果用于证明增量缓存的相对收益和工程可测性，不作为不同机器间的绝对性能承诺。"
    )

    add_h1(doc, "安装及使用")
    add_sub(doc, "1. 环境要求")
    add_bullets(
        doc,
        [
            "操作系统：Windows 10/11 或 macOS。当前本轮实测为 Windows + MSVC Debug 构建。",
            "构建工具：CMake 3.24 及以上、Ninja、支持 C++20 的编译器。",
            "桌面框架：Qt 6.9.x，Windows 推荐 msvc2022_64 Kit。",
            "语义分析可选依赖：LLVM/Clang 与 libclang；目标项目需要有效 compile_commands.json。",
            "AI 可选依赖：DeepSeek 官方 API 或 OpenAI 兼容代理端点，本地配置文件为 config/deepseek-ai.local.json，不纳入版本控制。",
        ],
    )
    add_sub(doc, "2. 构建过程")
    add_code_block(
        doc,
        "set SAVT_QT_ROOT=F:\\QT\\6.9.3\\msvc2022_64\n"
        "cmake --preset windows-msvc-qt-debug\n"
        "cmake --build --preset windows-msvc-qt-debug --parallel\n"
        "ctest --preset windows-msvc-qt-debug --output-on-failure",
    )
    add_para(
        doc,
        "如果当前 shell 中 cmake 或 MSVC 标准库不可用，应先进入 Visual Studio Developer Command Prompt x64，再执行构建。"
        "启用语义分析时需额外设置 SAVT_LLVM_ROOT，并以 -DSAVT_ENABLE_CLANG_TOOLING=ON 重新配置。"
    )
    add_sub(doc, "3. 使用流程")
    add_bullets(
        doc,
        [
            "启动 savt_desktop.exe，也可以在命令行附带待分析项目路径。",
            "在界面中选择项目文件夹，点击开始分析。",
            "查看系统上下文，确认项目入口、技术栈、精度状态和风险提示。",
            "进入架构全景图查看能力域，选择节点查看证据和上下游关系。",
            "对关键能力执行组件下钻，查看内部组件和协作关系。",
            "在配置页查看语义状态和 AST 预览；在报告页复制 Markdown 报告。",
            "如已配置 AI，手动触发节点、项目或报告解释；未配置 AI 时核心分析仍可使用。",
        ],
    )

    add_h1(doc, "项目总结")
    add_para(
        doc,
        "SAVT 当前已经形成较完整的工程化原型。它不是单纯的关系图绘制器，而是把源码扫描、语法/语义分析、架构模型、证据链、分层布局、"
        "增量缓存、桌面交互和受约束 AI 解释组织为一条可追溯链路。作品的主要价值在于帮助用户从源码事实出发理解复杂系统，"
        "并在精度不足时明确说明原因。"
    )
    add_para(
        doc,
        "开发过程中最大的困难来自 C++ 语义精度。工业级分析必须依赖 compile_commands.json、LLVM/Clang 环境、系统头文件解析和跨翻译单元合并。"
        "项目通过精度模式和阻断状态管理这一风险，选择诚实暴露条件不足，而不是静默退化。另一个困难来自 UI 迁移和交互复杂度，"
        "当前 greenfield 工作台已经成为主入口，但仍需治理历史 QML 资源，避免参赛打包时出现无关文件。"
    )
    add_para(
        doc,
        "后续工作优先级为：修复本轮失败测试；为参赛演示准备稳定样例和截图；补齐作品信息摘要、readme.txt、PPT、PDF、演示视频和源代码压缩包；"
        "准备可运行安装包；在具备 LLVM 的环境中跑通语义演示；整理第三方许可证清单；如使用 AI，则在作品信息摘要中特别说明工具名称、来源和合规边界。"
    )

    add_h1(doc, "参考文献")
    refs = [
        "中国大学生计算机设计大赛组织委员会. 关于举办“2026年（第19届）中国大学生计算机设计大赛”的通知及竞赛规程[EB/OL]. 2026-03-01. https://jwc.ustb.edu.cn/docs/2026-03/aec2284456ea4b00848d819edb8aa651.pdf",
        "ISO/IEC/IEEE. ISO/IEC/IEEE 42010:2022 Software, systems and enterprise - Architecture description[S]. Geneva: ISO, 2022. https://www.iso.org/standard/74393.html",
        "ISO/IEC. ISO/IEC 25010:2023 Systems and software engineering - Systems and software Quality Requirements and Evaluation (SQuaRE) - Product quality model[S]. Geneva: ISO, 2023. https://www.iso.org/standard/78176.html",
        "Clements P, Bachmann F, Bass L, et al. Documenting Software Architectures: Views and Beyond[M]. 2nd ed. Boston: Addison-Wesley, 2011.",
        "Ducasse S, Pollet D. Software Architecture Reconstruction: A Process-Oriented Taxonomy[J]. IEEE Transactions on Software Engineering, 2009, 35(4): 573-591. DOI: 10.1109/TSE.2009.19.",
        "Lattner C, Adve V. LLVM: A Compilation Framework for Lifelong Program Analysis and Transformation[C]//Proceedings of the 2004 International Symposium on Code Generation and Optimization. Palo Alto, 2004. DOI: 10.1109/CGO.2004.1281665.",
        "LLVM Project. LibTooling - Clang documentation[EB/OL]. https://clang.llvm.org/docs/LibTooling.html",
        "Tree-sitter contributors. Tree-sitter: An incremental parsing system for programming tools[EB/OL]. https://github.com/tree-sitter/tree-sitter",
        "The Qt Company. Qt 6.9 Documentation[EB/OL]. https://doc.qt.io/qt-6.9/",
        "Sugiyama K, Tagawa S, Toda M. Methods for Visual Understanding of Hierarchical System Structures[J]. IEEE Transactions on Systems, Man, and Cybernetics, 1981, SMC-11(2): 109-125. DOI: 10.1109/TSMC.1981.4308636.",
    ]
    for idx, ref in enumerate(refs, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.first_line_indent = Cm(-0.6)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(f"[{idx}] {ref}")
        set_run_font(run, size=9.5)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
