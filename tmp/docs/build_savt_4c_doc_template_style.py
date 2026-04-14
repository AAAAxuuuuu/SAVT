from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement


ROOT = Path(r"G:\SAVT")
DOWNLOADS = Path.home() / "Downloads"
TEMPLATE_DIR = next(
    d for d in DOWNLOADS.iterdir()
    if d.is_dir() and "2023" in d.name and any(f.name.startswith("01-3") for f in d.iterdir())
)
TEMPLATE = next(f for f in TEMPLATE_DIR.iterdir() if f.name.startswith("01-3") and f.suffix.lower() == ".docx")
OUTPUT_DIR = ROOT / "output" / "doc"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT = OUTPUT_DIR / "SAVT-软件应用与开发类作品设计和开发文档-模板样式保留版-2026-04-13.docx"


def set_paragraph_text_keep_runs(paragraph: Paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def insert_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.style = paragraph.style
    new_para.add_run(text)
    return new_para


def replace_instruction_with_section(paragraph: Paragraph, texts: list[str]) -> None:
    set_paragraph_text_keep_runs(paragraph, texts[0])
    anchor = paragraph
    for text in texts[1:]:
        anchor = insert_after(anchor, text)


def find_heading_indices(doc: Document) -> dict[str, int]:
    result = {}
    for idx, para in enumerate(doc.paragraphs):
        if para.text.strip() in {
            "需求分析",
            "概要设计",
            "详细设计",
            "测试报告",
            "安装及使用",
            "项目总结",
            "参考文献",
        }:
            result[para.text.strip()] = idx
    return result


sections = {
    "需求分析": [
        "SAVT 面向中大型软件项目的架构理解问题。随着项目规模增长，开发者面对的不再是单个源文件，而是跨目录、跨模块、跨语言、跨构建环境的复杂工程。传统文件树和全文检索能够定位局部代码，但难以回答系统入口在哪里、核心能力如何分布、模块之间为什么存在依赖、当前分析结论是否具备语义精度等问题。软件架构重构研究指出，架构并不总是显式存在于代码中，系统演进还会造成架构漂移，因此从源码中重建和校验架构是维护理解的重要辅助任务。",
        "本作品的目标用户包括软件工程课程团队、开源项目维护者、接手陌生仓库的新成员、架构评审人员和重构负责人。用户需要一种工具把源码事实转化为可交互、可解释、可追溯的架构视图，并在精度不足时明确告诉用户限制原因。SAVT 的核心需求包括项目扫描、源码解析、架构聚合、图形化展示和可信解释。",
        "竞品与对标分析：IDE 文件树和跳转工具适合局部阅读，但缺少全局架构视图；文档生成器适合生成接口说明，但难以表达精度状态和证据链；传统静态分析工具精度较高，但输出通常偏底层工程数据，交互展示不足；通用大模型自然语言解释能力强，但存在脱离证据产生幻觉和隐私边界不清的问题。SAVT 的差异化定位是把源码事实、精度状态、架构图、组件下钻、报告和受约束 AI 解释组织为一个统一闭环。",
        "需求边界方面，SAVT 不承诺在缺少编译数据库时完成工业级 C++ 语义分析；在没有网络和 API 配置时，AI 解释功能可以不可用，但项目扫描、语法分析、架构图、组件下钻、AST 预览和报告生成仍应独立运行。",
    ],
    "概要设计": [
        "SAVT 采用分层架构，核心原则是分析事实与界面渲染解耦。底层负责项目扫描和源码解析，中间层负责统一图模型、能力图、组件图和证据包，上层负责布局计算、Qt 服务编排、QML 场景映射和桌面交互展示。该设计与 ISO/IEC/IEEE 42010 中关于架构描述、视图和模型的思想一致，即一个架构描述应围绕利益相关者关切组织多个视图，而不是只给出单一图形。",
        "总体调用关系：QML Desktop Workbench -> AnalysisController/AnalysisOrchestrator -> IncrementalAnalysisPipeline -> CppProjectAnalyzer -> SyntaxProjectAnalyzer 或 SemanticProjectAnalyzer -> AnalysisReport -> ArchitectureOverview/CapabilityGraph/ComponentGraph -> LayeredGraphLayout -> SceneMapper -> L1-L4 视图、报告、AST 预览和 AI 解释。",
        "核心模块划分：apps/desktop 负责桌面入口和 QML 工作台；src/core 定义 AnalysisReport、ArchitectureOverview、CapabilityGraph、ComponentGraph 和证据包；src/parser 封装 Tree-sitter C++ 解析；src/analyzer 负责项目扫描、符号关系抽取和语法/语义精度模式；src/layout 负责能力图、组件图和模块图布局；src/ui 负责异步分析、缓存、报告、AST 预览、语义状态、AI 服务和 scene 映射；src/ai 负责 DeepSeek/OpenAI 兼容配置、请求构造和结构化响应解析；tests 负责后端、快照、AI、UI 服务层、性能基线和 UTF-8 检查。",
        "系统数据流从用户选择项目目录开始。分析器先生成扫描清单和元数据指纹，再根据精度模式选择 Tree-sitter 语法路径或 libclang 语义路径。完成分析后，结果被聚合为架构总览和能力图，再进入布局模块生成可视化坐标，最终由 QML 工作台展示。项目不使用业务数据库，主要状态来自内存模型、项目级 JSON 配置和可选本地 AI 配置文件。",
    ],
    "详细设计": [
        "一、精度模式与语义诊断。SAVT 定义 SyntaxOnly、SemanticPreferred 和 SemanticRequired 三种分析模式。SyntaxOnly 只执行语法级分析；SemanticPreferred 优先尝试语义分析，失败时可回退到语法分析；SemanticRequired 必须进入语义分析，否则直接生成阻断报告。阻断原因包括 missing_compile_commands、backend_unavailable、llvm_not_found、llvm_headers_missing、libclang_not_found、system_headers_unresolved 和 translation_unit_parse_failed 等。",
        "二、核心算法流程。若精度模式请求语义分析，系统先查找 compile_commands.json，再检查 LLVM/libclang 可用性；当编译数据库和语义后端可用时运行语义分析并返回语义报告；当语义必需模式下条件不满足时返回带具体原因的阻断报告；其他情况下运行 Tree-sitter 语法分析并返回 syntax 或 syntax_fallback 报告。该流程保证“分析结果”和“精度声明”一致。",
        "三、图模型与证据包。底层 AnalysisReport 保存符号节点、关系边、解析文件数、语义状态码和诊断信息。CapabilityGraph 将模块组织为入口、能力和基础设施节点，并在每个节点上保留事实、规则、结论、来源文件、符号、模块、关系、置信标签和置信原因。该证据包机制使评委或开发者可以追问“为什么存在这个节点”以及“这条边来自什么证据”，避免架构图变成不可审查的静态图片。",
        "四、增量分析与缓存。增量管线把一次分析拆分为扫描、解析、聚合、布局四层缓存。扫描层缓存项目源码清单和元数据指纹；解析层缓存 AnalysisReport；聚合层缓存 ArchitectureOverview、CapabilityGraph 和规则报告；布局层缓存 L2/L3 scene layout 和组件图。源码、项目配置、compile_commands.json 或精度模式变化时，对应层及其下游层会失效重算。",
        "五、人机界面与典型流程。当前桌面主界面由 Qt Quick/QML 实现，主入口为 apps/desktop/qml/Main.qml 和 greenfield 工作台。用户可以选择项目目录、启动分析、停止分析、查看语义状态、在架构全景图中缩放和平移、选择节点、聚焦上下游关系、下钻组件图、查看 AST 预览，并复制 Markdown 报告。最终参赛版建议补充实际界面截图和 10 分钟演示视频，本稿不伪造截图。",
        "六、AI 辅助解释边界。AI 功能是可选辅助层。系统只在用户手动触发时读取受限证据包，构造 DeepSeek 或 OpenAI 兼容 Chat Completions 请求。内置规则要求模型只处理 SAVT 架构阅读任务，只使用提供的证据，不臆造文件、模块、符号或业务含义，并以结构化 JSON 返回。无网络或无本地密钥配置时，不影响核心静态分析功能。",
    ],
    "测试报告": [
        "本节数据为 2026年4月13日 在当前 Windows 工作区实测结果。构建命令在 Visual Studio 2022 Developer Command Prompt x64 环境下执行，目标 savt_backend_tests、savt_snapshot_tests、savt_ai_tests、savt_ui_tests、savt_perf_baseline 已构建成功。随后执行 ctest 全量测试。",
        "基础统计：当前 Git 提交为 87828d6f6dd974ed92f8fa93525debd65af4ad7a；统计范围包括 src、apps/desktop、tests/cpp、config、scripts、samples、docs，排除 build、third_party 和 docs/competition。扫描得到 179 个文件，其中 QML 68 个、C++ 源文件 35 个、头文件 29 个、Markdown 25 个。主要代码行数为 C/C++ 相关 63 文件 19165 行、QML/JS 70 文件 12200 行、C++ 测试 5 文件 2674 行。精度样例 tests/fixtures/precision 为 21 个，黄金结果 tests/golden/precision 为 21 个。使用 rg 扫描 TODO/FIXME/HACK/XXX/@deprecated 未发现匹配项。",
        "ctest 全量结果：6 个测试项中 3 个通过、3 个失败。通过项为 savt_ai_tests、savt_perf_smoke、savt_utf8_check。失败项为 savt_backend_tests、savt_snapshot_tests、savt_ui_tests。savt_backend_tests 的失败断言为 entry overview modules should become entry components in the L3 graph；savt_snapshot_tests 出现多个 fixture 与 golden 快照不一致，说明当前实现已改变输出，需要审查并更新基线或修复回归；savt_ui_tests 的失败断言为 system context cards should publish the expected summary cards。",
        "因此，本文档不将当前仓库描述为“全部测试通过”的最终交付状态。更准确的结论是：项目可构建，AI、性能 smoke 和编码检查通过；后端 L3 组件映射、系统上下文卡片和快照基线存在待修复或待重新审定项。正式参赛前必须修复失败测试，或对 golden 基线变化给出代码审查记录。",
        "性能 smoke 实测结果：cpp_vendored_dependency_app 发现 6 个文件、解析 6 个文件、能力节点/边为 2/1、冷启动总耗时 19.92 ms、热缓存总耗时 3.02 ms、四层缓存命中、峰值内存 18.23 MB；qml_dashboard_workspace 发现 8 个文件、解析 5 个文件、能力节点/边为 5/2、冷启动 18.33 ms、热缓存 6.17 ms、四层缓存命中、峰值内存 18.47 MB；synthetic_quick_2000 发现 2001 个文件、解析 2001 个文件、能力节点/边为 9/0、冷启动 2350.33 ms、热缓存 469.15 ms、四层缓存命中、峰值内存 77.93 MB。",
        "上述性能数据来自 savt_perf_smoke 的本轮实测输出。由于测试机器负载、Debug 构建、磁盘状态和缓存状态会影响绝对耗时，该结果用于证明增量缓存的相对收益和工程可测性，不作为不同机器间的绝对性能承诺。",
    ],
    "安装及使用": [
        "环境要求：Windows 10/11 或 macOS；CMake 3.24 及以上；Ninja；支持 C++20 的编译器；Qt 6.9.x，Windows 推荐 msvc2022_64 Kit。语义分析为可选能力，需要 LLVM/Clang 与 libclang，且目标项目需要有效 compile_commands.json。AI 为可选能力，需要 DeepSeek 官方 API 或 OpenAI 兼容代理端点，本地配置文件为 config/deepseek-ai.local.json，不纳入版本控制。",
        "推荐构建流程：设置 SAVT_QT_ROOT 指向 Qt Kit 根目录，执行 cmake --preset windows-msvc-qt-debug，再执行 cmake --build --preset windows-msvc-qt-debug --parallel，最后执行 ctest --preset windows-msvc-qt-debug --output-on-failure。如果当前 shell 中 cmake 或 MSVC 标准库不可用，应先进入 Visual Studio Developer Command Prompt x64，再执行构建。启用语义分析时需额外设置 SAVT_LLVM_ROOT，并以 -DSAVT_ENABLE_CLANG_TOOLING=ON 重新配置。",
        "典型使用流程：启动 savt_desktop.exe，也可以在命令行附带待分析项目路径；在界面中选择项目文件夹，点击开始分析；查看系统上下文，确认项目入口、技术栈、精度状态和风险提示；进入架构全景图查看能力域，选择节点查看证据和上下游关系；对关键能力执行组件下钻，查看内部组件和协作关系；在配置页查看语义状态和 AST 预览；在报告页复制 Markdown 报告；如已配置 AI，可手动触发节点、项目或报告解释；未配置 AI 时核心分析仍可使用。",
    ],
    "项目总结": [
        "SAVT 当前已经形成较完整的工程化原型。它不是单纯的关系图绘制器，而是把源码扫描、语法/语义分析、架构模型、证据链、分层布局、增量缓存、桌面交互和受约束 AI 解释组织为一条可追溯链路。作品的主要价值在于帮助用户从源码事实出发理解复杂系统，并在精度不足时明确说明原因。",
        "开发过程中最大的困难来自 C++ 语义精度。工业级分析必须依赖 compile_commands.json、LLVM/Clang 环境、系统头文件解析和跨翻译单元合并。项目通过精度模式和阻断状态管理这一风险，选择诚实暴露条件不足，而不是静默退化。另一个困难来自 UI 迁移和交互复杂度，当前 greenfield 工作台已经成为主入口，但仍需治理历史 QML 资源，避免参赛打包时出现无关文件。",
        "后续工作优先级为：修复本轮失败测试；为参赛演示准备稳定样例和截图；补齐作品信息摘要、readme.txt、PPT、PDF、演示视频和源代码压缩包；准备可运行安装包；在具备 LLVM 的环境中跑通语义演示；整理第三方许可证清单；如使用 AI，则在作品信息摘要中特别说明工具名称、来源和合规边界。",
    ],
    "参考文献": [
        "[1] 中国大学生计算机设计大赛组织委员会. 关于举办“2026年（第19届）中国大学生计算机设计大赛”的通知及竞赛规程[EB/OL]. 2026-03-01. https://jwc.ustb.edu.cn/docs/2026-03/aec2284456ea4b00848d819edb8aa651.pdf",
        "[2] ISO/IEC/IEEE. ISO/IEC/IEEE 42010:2022 Software, systems and enterprise - Architecture description[S]. Geneva: ISO, 2022. https://www.iso.org/standard/74393.html",
        "[3] ISO/IEC. ISO/IEC 25010:2023 Systems and software engineering - Systems and software Quality Requirements and Evaluation (SQuaRE) - Product quality model[S]. Geneva: ISO, 2023. https://www.iso.org/standard/78176.html",
        "[4] Clements P, Bachmann F, Bass L, et al. Documenting Software Architectures: Views and Beyond[M]. 2nd ed. Boston: Addison-Wesley, 2011.",
        "[5] Ducasse S, Pollet D. Software Architecture Reconstruction: A Process-Oriented Taxonomy[J]. IEEE Transactions on Software Engineering, 2009, 35(4): 573-591. DOI: 10.1109/TSE.2009.19.",
        "[6] Lattner C, Adve V. LLVM: A Compilation Framework for Lifelong Program Analysis and Transformation[C]//Proceedings of the 2004 International Symposium on Code Generation and Optimization. Palo Alto, 2004. DOI: 10.1109/CGO.2004.1281665.",
        "[7] LLVM Project. LibTooling - Clang documentation[EB/OL]. https://clang.llvm.org/docs/LibTooling.html",
        "[8] Tree-sitter contributors. Tree-sitter: An incremental parsing system for programming tools[EB/OL]. https://github.com/tree-sitter/tree-sitter",
        "[9] The Qt Company. Qt 6.9 Documentation[EB/OL]. https://doc.qt.io/qt-6.9/",
        "[10] Sugiyama K, Tagawa S, Toda M. Methods for Visual Understanding of Hierarchical System Structures[J]. IEEE Transactions on Systems, Man, and Cybernetics, 1981, SMC-11(2): 109-125. DOI: 10.1109/TSMC.1981.4308636.",
    ],
}


def build() -> None:
    doc = Document(TEMPLATE)

    set_paragraph_text_keep_runs(doc.paragraphs[4], "软件应用与开发类作品设计和开发文档")
    set_paragraph_text_keep_runs(doc.paragraphs[5], "作品编号：待报名系统生成后填写")
    set_paragraph_text_keep_runs(doc.paragraphs[6], "作品名称：SAVT 软件架构分析与可视化系统")
    set_paragraph_text_keep_runs(doc.paragraphs[7], "作　　者：待参赛团队确认后填写")
    set_paragraph_text_keep_runs(doc.paragraphs[8], "版本编号：v0.1.0-competition-draft, Git 87828d6")
    set_paragraph_text_keep_runs(doc.paragraphs[9], "填写日期：2026年4月13日")

    section_order = [
        "参考文献",
        "项目总结",
        "安装及使用",
        "测试报告",
        "详细设计",
        "概要设计",
        "需求分析",
    ]
    for heading in section_order:
        headings = find_heading_indices(doc)
        texts = sections[heading]
        idx = headings[heading]
        if idx + 1 >= len(doc.paragraphs):
            continue
        replace_instruction_with_section(doc.paragraphs[idx + 1], texts)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
